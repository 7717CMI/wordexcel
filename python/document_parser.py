import os
import re
from typing import Dict, Any, List
from docx import Document
import docx2txt
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_word(file_path: str) -> str:
    """Extract text from Word document using appropriate library based on file extension"""
    try:
        logger.info('Starting Word document text extraction...')
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_size = os.path.getsize(file_path)
        logger.info(f'File read successfully, size: {file_size} bytes')
        
        # Determine file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        logger.info(f'File extension: {file_extension}')
        
        extracted_text = ""
        
        if file_extension == '.docx':
            # Use python-docx for .docx files (Office Open XML format)
            logger.info('Processing .docx file with python-docx')
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            extracted_text = '\n'.join(text_parts)
            
        elif file_extension == '.doc':
            # Use multiple approaches for .doc files (Microsoft Word 97-2003 format)
            logger.info('Processing .doc file with multiple approaches')
            extracted_text = ""
            
            # Method 1: Try docx2txt first
            try:
                logger.info('Trying docx2txt...')
                extracted_text = docx2txt.process(file_path)
                if extracted_text and len(extracted_text.strip()) >= 10:
                    logger.info('✅ docx2txt succeeded')
                else:
                    raise Exception("Extracted text is too short or empty")
            except Exception as docx2txt_error:
                logger.warning(f'docx2txt failed: {docx2txt_error}')
                
                # Method 2: Try python-docx as fallback (might work for some .doc files)
                try:
                    logger.info('Trying python-docx as fallback...')
                    doc = Document(file_path)
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_parts.append(' | '.join(row_text))
                    extracted_text = '\n'.join(text_parts)
                    
                    if extracted_text and len(extracted_text.strip()) >= 10:
                        logger.info('✅ python-docx fallback succeeded')
                    else:
                        raise Exception("Extracted text is too short or empty")
                        
                except Exception as docx_error:
                    logger.warning(f'python-docx fallback failed: {docx_error}')
                    
                    # Method 3: Try using subprocess with antiword (if available)
                    try:
                        logger.info('Trying antiword...')
                        import subprocess
                        result = subprocess.run(['antiword', file_path], 
                                              capture_output=True, text=True, timeout=30)
                        if result.returncode == 0 and result.stdout.strip():
                            extracted_text = result.stdout
                            if len(extracted_text.strip()) >= 10:
                                logger.info('✅ antiword succeeded')
                            else:
                                raise Exception("antiword output too short")
                        else:
                            raise Exception(f"antiword failed: {result.stderr}")
                    except Exception as antiword_error:
                        logger.warning(f'antiword failed: {antiword_error}')
                        
                        # Method 4: Try reading as binary and extract text manually
                        try:
                            logger.info('Trying binary text extraction...')
                            with open(file_path, 'rb') as f:
                                content = f.read()
                            
                            # Look for text content in binary data
                            import re
                            # Try to find readable text patterns
                            text_patterns = [
                                rb'[\x20-\x7E]{10,}',  # Printable ASCII characters
                                rb'[\x09\x0A\x0D\x20-\x7E]{10,}',  # Including whitespace
                            ]
                            
                            for pattern in text_patterns:
                                matches = re.findall(pattern, content)
                                if matches:
                                    # Decode and join matches
                                    text_parts = []
                                    for match in matches:
                                        try:
                                            decoded = match.decode('utf-8', errors='ignore')
                                            if len(decoded.strip()) > 5:
                                                text_parts.append(decoded.strip())
                                        except:
                                            continue
                                    
                                    if text_parts:
                                        extracted_text = '\n'.join(text_parts)
                                        if len(extracted_text.strip()) >= 10:
                                            logger.info('✅ Binary text extraction succeeded')
                                            break
                            
                            if not extracted_text or len(extracted_text.strip()) < 10:
                                raise Exception("Binary extraction failed")
                                
                        except Exception as binary_error:
                            logger.error(f'Binary extraction failed: {binary_error}')
                            raise Exception(f'Failed to read .doc file with all methods: docx2txt, python-docx, antiword, and binary extraction')
            
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}. Only .doc and .docx files are supported.")
        
        logger.info(f'Text extracted successfully, length: {len(extracted_text)}')
        logger.info(f'Text preview (first 200 chars): {extracted_text[:200]}')
        
        # Basic cleanup
        cleaned_text = extracted_text.replace('\r\n', '\n').replace('\r', '\n').strip()
        logger.info(f'Cleaned text length: {len(cleaned_text)}')
        
        return cleaned_text
        
    except Exception as error:
        logger.error(f'Error reading Word document: {error}')
        raise Exception(f'Failed to read Word document: {str(error)}')

def preprocess_text(text: str) -> str:
    """Basic text preprocessing for AI analysis"""
    return text.replace('\r\n', '\n').replace('\r', '\n').replace('\n\\s*\n', '\n').replace('\\s+', ' ').strip()

def normalize_extracted_data(data: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize extracted data to handle common variations"""
    try:
        # Handle common field name variations
        if 'market' in data:
            market = data['market']
            if 'marketName' in market and 'market_name' not in market:
                market['market_name'] = market['marketName']
            if 'baseYear' in market and 'base_year' not in market:
                market['base_year'] = market['baseYear']
            if 'startYear' in market and 'start_year' not in market:
                market['start_year'] = market['startYear']
            if 'endYear' in market and 'end_year' not in market:
                market['end_year'] = market['endYear']
            if 'sizeBase' in market and 'size_base_raw' not in market:
                market['size_base_raw'] = market['sizeBase']
            if 'sizeForecast' in market and 'size_forecast_raw' not in market:
                market['size_forecast_raw'] = market['sizeForecast']
            if 'cagr' in market and 'cagr_percent_display' not in market:
                market['cagr_percent_display'] = market['cagr']
            if 'currency' in market and 'currency_unit' not in market:
                market['currency_unit'] = market['currency']

        # Handle common segment field name variations
        if 'segments' in data:
            segments = data['segments']
            segment_keys = list(segments.keys())
            
            logger.info('=== NORMALIZATION DEBUG ===')
            logger.info(f'Original segment keys: {segment_keys}')
            logger.info(f'Original segments structure: {segments}')
            
            if segment_keys:
                # Create normalized structure with cat_1, cat_2, cat_3, cat_4, cat_5
                normalized_segments = {}
                
                # Map up to 5 categories to cat_1, cat_2, cat_3, cat_4, cat_5
                for index, key in enumerate(segment_keys[:5]):
                    category_number = index + 1
                    normalized_key = f'cat_{category_number}'
                    
                    if key in segments and isinstance(segments[key], dict):
                        normalized_segments[normalized_key] = {
                            'header': segments[key].get('header', key),
                            'items': segments[key].get('items', []),
                            'shares': segments[key].get('shares', [])
                        }
                        logger.info(f'Normalized {key} -> {normalized_key}: {normalized_segments[normalized_key]}')
                
                # Initialize empty categories for any missing ones (up to 5 total)
                for i in range(1, 6):
                    cat_key = f'cat_{i}'
                    if cat_key not in normalized_segments:
                        normalized_segments[cat_key] = {
                            'header': '',
                            'items': [],
                            'shares': []
                        }
                        logger.info(f'Initialized empty {cat_key}')
                
                logger.info(f'Final normalized segments: {normalized_segments}')
                
                # Replace the original segments with normalized structure
                data['segments'] = normalized_segments

        # Handle common players field name variations
        if 'players' in data:
            players = data['players']
            if 'companies' in players and 'players' not in players:
                players['players'] = players['companies']
            if 'keyPlayers' in players and 'players' not in players:
                players['players'] = players['keyPlayers']

        return data
    except Exception as error:
        logger.error(f'Error normalizing data: {error}')
        return data

def validate_extracted_data(data: Dict[str, Any]) -> bool:
    """Validate extracted data structure"""
    try:
        logger.info('=== VALIDATION DEBUG ===')
        logger.info(f'Raw data to validate: {data}')
        logger.info(f'Data type: {type(data)}')
        
        # Check if data is a dictionary
        if not isinstance(data, dict):
            logger.error(f'Data is not a dictionary: {type(data)}')
            return False
        
        # Check if required fields exist
        required_keys = ['market', 'segments', 'players']
        missing_keys = [key for key in required_keys if key not in data]
        
        if missing_keys:
            logger.error(f'Missing required top-level fields: {missing_keys}')
            logger.error(f'Available keys: {list(data.keys())}')
            return False

        # Validate market data
        market = data['market']
        if not isinstance(market, dict):
            logger.error(f'Market data is not a dictionary: {type(market)}')
            return False
            
        required_market_fields = ['market_name', 'base_year', 'start_year', 'end_year']
        missing_market_fields = [field for field in required_market_fields if field not in market]
        
        if missing_market_fields:
            logger.error(f'Missing required market fields: {missing_market_fields}')
            logger.error(f'Available market fields: {list(market.keys())}')
            return False

        # Validate segments
        segments = data['segments']
        if not isinstance(segments, dict):
            logger.error(f'Segments data is not a dictionary: {type(segments)}')
            return False
            
        logger.info(f'Available segment categories: {list(segments.keys())}')
        
        # Check if we have at least one valid segment category
        has_valid_segment = False
        for key, segment in segments.items():
            if (isinstance(segment, dict) and 
                'header' in segment and 
                'items' in segment and
                isinstance(segment['header'], str) and 
                isinstance(segment['items'], list)):
                has_valid_segment = True
                logger.info(f'Valid segment found: {key} with header "{segment["header"]}" and {len(segment["items"])} items')
                break
        
        if not has_valid_segment:
            logger.error('No valid segment categories found')
            logger.error(f'Segment data structure: {segments}')
            return False

        # Validate players
        players = data['players']
        if not isinstance(players, dict):
            logger.error(f'Players data is not a dictionary: {type(players)}')
            return False
            
        if 'players' not in players or not isinstance(players['players'], list):
            logger.error(f'Invalid players data: {players}')
            return False
        
        logger.info('✅ Validation passed successfully')
        return True
        
    except Exception as error:
        logger.error(f'Data validation error: {error}')
        logger.error(f'Data that caused error: {data}')
        return False

def calculate_confidence(data: Dict[str, Any]) -> int:
    """Calculate confidence score based on data completeness"""
    try:
        score = 0
        total = 0

        # Market data completeness
        market_fields = [
            'market_name', 'base_year', 'start_year', 'end_year',
            'size_base_raw', 'size_forecast_raw', 'cagr_percent_display', 'currency_unit'
        ]
        
        for field in market_fields:
            total += 1
            if field in data['market'] and data['market'][field]:
                score += 1

        # Segments completeness
        segment_categories = ['cat_1', 'cat_2', 'cat_3']
        for cat in segment_categories:
            total += 1
            if (cat in data['segments'] and 
                data['segments'][cat].get('header') and 
                data['segments'][cat].get('items')):
                score += 1

        # Players completeness
        total += 1
        if data['players'].get('players'):
            score += 1

        return round((score / total) * 100)
    except Exception as error:
        logger.error(f'Error calculating confidence: {error}')
        return 0
